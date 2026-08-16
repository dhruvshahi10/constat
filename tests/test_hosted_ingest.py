"""Eval (a): upload ingestion round-trip.

An upload that persists must be ingestible by the exact reader the pipeline
uses; attestation is the only path to citability; failures are clean.
"""
import io
from pathlib import Path

import pytest

from pramana.evidence import parse_source
from pramana.server import extract
from pramana.server.multipart import MultipartError, parse

GOOD_META = {
    "title": "Access Control Policy", "type": "policy", "version": "2.1",
    "owner": "security@example.com", "effective_date": "2026-01-01",
    "expiry_date": "2027-01-01", "topics": "access, mfa", "attested": "on",
}

BODY = ("All workforce access to production systems requires multi-factor "
        "authentication.\n\nAccess rights are reviewed quarterly.")


def _mk_docx() -> bytes:
    import docx
    d = docx.Document()
    d.add_paragraph("All workforce access requires multi-factor authentication.")
    d.add_paragraph("Access reviews are quarterly.")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _mk_pdf() -> bytes:
    """Minimal one-page PDF with real extractable text."""
    text = "MFA is enforced for all production access. Reviews are quarterly."
    stream = f"BT /F1 12 Tf 50 700 Td ({text}) Tj ET".encode()
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R "
        b"/Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = io.BytesIO()
    out.write(b"%PDF-1.4\n")
    offsets = []
    for i, obj in enumerate(objs, start=1):
        offsets.append(out.tell())
        out.write(f"{i} 0 obj\n".encode() + obj + b"\nendobj\n")
    xref = out.tell()
    out.write(f"xref\n0 {len(objs)+1}\n0000000000 65535 f \n".encode())
    for off in offsets:
        out.write(f"{off:010d} 00000 n \n".encode())
    out.write(b"trailer\n<< /Size " + str(len(objs) + 1).encode() +
              b" /Root 1 0 R >>\nstartxref\n" + str(xref).encode() + b"\n%%EOF")
    return out.getvalue()


# --- extraction per format ---------------------------------------------------
def test_md_txt_extract():
    assert "multi-factor" in extract.extract_text("p.txt", BODY.encode())
    md = "---\nsource_id: EVIL\napproval_status: approved\n---\n" + BODY
    text = extract.extract_text("p.md", md.encode())
    assert "EVIL" not in text  # uploaded frontmatter is never trusted


def test_docx_extract():
    assert "multi-factor" in extract.extract_text("p.docx", _mk_docx())


def test_pdf_extract():
    assert "MFA is enforced" in extract.extract_text("p.pdf", _mk_pdf())


def test_corrupt_pdf_fails_clean():
    with pytest.raises(extract.ExtractError, match="Could not read|No extractable"):
        extract.extract_text("bad.pdf", b"%PDF-1.4 not really a pdf")


def test_unsupported_ext_rejected():
    with pytest.raises(extract.ExtractError, match="Unsupported"):
        extract.extract_text("evil.exe", b"x")


# --- synthesis round-trip ----------------------------------------------------
def test_synthesis_roundtrip(tmp_path):
    meta = extract.validate_meta(GOOD_META)
    sid, sha, text = extract.synthesize(meta, BODY, "acme")
    p = tmp_path / f"{sid}.md"
    p.write_text(text, encoding="utf-8")
    src = parse_source(p, "acme")
    assert src.source_id == sid
    assert src.is_approved()
    assert src.type == "policy"
    assert len(sha) == 64
    assert sid.startswith("POL-")


def test_unattested_is_not_citable(tmp_path):
    meta = extract.validate_meta({**GOOD_META, "attested": ""})
    sid, _, text = extract.synthesize(meta, BODY, "acme")
    p = tmp_path / f"{sid}.md"
    p.write_text(text, encoding="utf-8")
    assert parse_source(p, "acme").is_approved() is False


@pytest.mark.parametrize("field,val,msg", [
    ("title", "x", "Title"),
    ("type", "meme", "Type"),
    ("owner", "not-an-email", "email"),
    ("effective_date", "01/01/2026", "YYYY-MM-DD"),
    ("expiry_date", "2025-01-01", "after"),
])
def test_meta_validation(field, val, msg):
    with pytest.raises(extract.ExtractError, match=msg):
        extract.validate_meta({**GOOD_META, field: val})


# --- multipart parser --------------------------------------------------------
def _multipart_body(fields: dict[str, str], file: tuple[str, bytes]) -> tuple[bytes, str]:
    b = "xBOUNDx"
    out = io.BytesIO()
    for k, v in fields.items():
        out.write(f'--{b}\r\nContent-Disposition: form-data; name="{k}"\r\n\r\n{v}\r\n'.encode())
    fname, data = file
    out.write(f'--{b}\r\nContent-Disposition: form-data; name="file"; '
              f'filename="{fname}"\r\nContent-Type: application/pdf\r\n\r\n'.encode())
    out.write(data)
    out.write(f"\r\n--{b}--\r\n".encode())
    return out.getvalue(), f"multipart/form-data; boundary={b}"


def test_multipart_binary_survives():
    payload = bytes(range(256)) * 3 + b"\r\n--fake--\r\n\x00binary"
    body, ctype = _multipart_body({"title": "T"}, ("f.pdf", payload))
    parts = parse(body, ctype)
    assert parts["file"] == ("f.pdf", payload)
    assert parts["title"][1] == b"T"


def test_multipart_rejects_garbage():
    with pytest.raises(MultipartError):
        parse(b"no boundary here", "multipart/form-data; boundary=zzz")
    with pytest.raises(MultipartError):
        parse(b"x", "text/plain")


# --- regression: extraction must be bounded BEFORE it consumes memory --------
# A re-audit built two bombs that pass the 5MB wire gate: a 0.41MB DOCX
# inflating to 400MB (826MB peak RSS) and a 0.44MB DOCX of 1.2M paragraphs
# (68s CPU, 1.1GB RSS) before the post-extraction character check fired.
# Render starter is 512MB, so either is a single-request OOM that takes the run
# worker and every in-flight run with it.

def _zip_bomb(uncompressed_mb: int) -> bytes:
    import io, zipfile
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", "<x/>")
        z.writestr("word/document.xml", b"A" * (uncompressed_mb * 1024 * 1024))
    return buf.getvalue()


def test_docx_decompression_bomb_refused_from_the_zip_directory():
    data = _zip_bomb(400)
    assert len(data) < 5 * 1024 * 1024, "fixture must pass the wire-size gate"
    with pytest.raises(extract.ExtractError, match="expands to more"):
        extract.extract_text("bomb.docx", data)


def test_bomb_is_refused_without_parsing_it():
    """The refusal must come from the zip directory, not from parsing."""
    import time
    data = _zip_bomb(400)
    t0 = time.perf_counter()
    with pytest.raises(extract.ExtractError):
        extract.extract_text("bomb.docx", data)
    assert time.perf_counter() - t0 < 1.0, "refused too slowly to be pre-parse"


def test_ordinary_docx_still_ingests():
    """The bomb guard must not refuse a legitimate document."""
    text = extract.extract_text("policy.docx", _mk_docx())
    assert "multi-factor" in text
